#!/usr/bin/env python3
"""
Script to create a professional DOCX document from markdown files
"""

import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
import re

def add_hyperlink(paragraph, url, text):
    """Add a hyperlink to a paragraph"""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Add hyperlink styling
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)
    
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    
    paragraph._p.append(hyperlink)
    return hyperlink

def create_professional_docx():
    # Create a new document
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Add title page
    title = doc.add_heading('Peach Frontend', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('API Integrations & RPC Endpoints Documentation', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add some spacing
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Add document info
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_para.add_run('Technical Documentation\n')
    info_para.add_run('Version 1.0\n')
    info_para.add_run('Generated from codebase analysis')
    
    # Add page break
    doc.add_page_break()
    
    # Read the API documentation
    api_doc_path = '/Users/macbook/project/peach-fe/API_INTEGRATIONS_DOCUMENTATION.md'
    rpc_doc_path = '/Users/macbook/project/peach-fe/RPC_ENDPOINTS_DOCUMENTATION.md'
    
    def process_markdown_content(file_path, doc):
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        lines = content.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i].strip()
            
            if not line:
                doc.add_paragraph()
                i += 1
                continue
            
            # Handle headers
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('#### '):
                doc.add_heading(line[5:], level=4)
            elif line.startswith('##### '):
                doc.add_heading(line[6:], level=5)
            
            # Handle code blocks
            elif line.startswith('```'):
                code_lines = []
                i += 1
                while i < len(lines) and not lines[i].strip().startswith('```'):
                    code_lines.append(lines[i])
                    i += 1
                
                if code_lines:
                    code_para = doc.add_paragraph()
                    code_run = code_para.add_run('\n'.join(code_lines))
                    code_run.font.name = 'Courier New'
                    code_run.font.size = Pt(9)
                    
                    # Add background color (light gray)
                    from docx.oxml.shared import OxmlElement
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'), 'F5F5F5')
                    code_run._element.get_or_add_rPr().append(shd)
            
            # Handle bullet points
            elif line.startswith('- '):
                para = doc.add_paragraph(line[2:], style='List Bullet')
            
            # Handle numbered lists
            elif re.match(r'^\d+\. ', line):
                para = doc.add_paragraph(line[3:], style='List Number')
            
            # Handle horizontal rules
            elif line.startswith('---'):
                doc.add_paragraph()
                para = doc.add_paragraph()
                para.add_run('_' * 50)
                doc.add_paragraph()
            
            # Handle bold text with **Location**:
            elif '**Location**:' in line:
                para = doc.add_paragraph()
                parts = line.split('**Location**:')
                if len(parts) == 2:
                    para.add_run('Location: ').bold = True
                    para.add_run(parts[1].strip())
            
            # Handle other bold patterns
            elif '**' in line:
                para = doc.add_paragraph()
                # Simple bold text handling
                parts = re.split(r'\*\*(.*?)\*\*', line)
                for j, part in enumerate(parts):
                    if j % 2 == 0:
                        para.add_run(part)
                    else:
                        para.add_run(part).bold = True
            
            # Regular paragraphs
            else:
                doc.add_paragraph(line)
            
            i += 1
    
    # Process API documentation
    process_markdown_content(api_doc_path, doc)
    
    # Add page break before RPC documentation
    doc.add_page_break()
    
    # Process RPC documentation
    process_markdown_content(rpc_doc_path, doc)
    
    # Save the document
    output_path = '/Users/macbook/project/peach-fe/Peach_Frontend_API_RPC_Documentation.docx'
    doc.save(output_path)
    
    print(f"Professional DOCX document created successfully: {output_path}")
    print("Document features:")
    print("- Professional formatting with proper margins")
    print("- Title page with document information")
    print("- Structured headings and subheadings")
    print("- Code blocks with monospace font and background")
    print("- Bullet points and numbered lists")
    print("- Bold text formatting")
    print("- Compatible with Microsoft Word, Google Docs, and other word processors")

if __name__ == '__main__':
    create_professional_docx()